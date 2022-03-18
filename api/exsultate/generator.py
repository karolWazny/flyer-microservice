import os
import time
import re
from .exsultate import *
from docx import Document
from pathlib import Path
from io import BytesIO
import json

class Generator:
    def __init__(self):
        self.path = Path(os.path.dirname(os.path.realpath(__file__)))
        self.configuration = None
        self.document = None
        self.previous_indented = None

    def load_configuration(self, filename="config.json"):
        if not os.path.isfile(self.path.joinpath(Path(filename))):
            self.generate_config_file(self.path.joinpath(Path(filename)))
        self.configuration = read_json(self.path.joinpath(Path(filename)))

    @staticmethod
    def get_default_config():
        return {"style_mappings": { "number": "number", "verse": "verse", "chorus": "chorus", "title": "title-itself",
                                    "prayer":"prayer"},
                "template": "template.docx",
                "prayer":"prayer.json"}

    def generate_config_file(self, filename):
        write_json(Generator.get_default_config(), filename)

    def generate(self, songbook, target_stream=None):
        if target_stream is None:
            target_stream = BytesIO()
        self.document = Document(self.path.joinpath(Path(self.configuration['template'])))
        for song in songbook.songs:
            self.add_song_to_document(song)
        self.add_prayer()
        self.document.save(target_stream)
        return target_stream

    def add_prayer(self):
        style_name = self.configuration['style_mappings']['prayer']
        try:
            file = open(self.path.joinpath(Path(self.configuration['prayer'])), 'r')
            content = json.load(file)['content']
            self.document.add_paragraph(content, style=style_name)
        except FileNotFoundError:
            pass

    def add_song_to_document(self, song):
        self.previous_indented = True
        self.document.add_paragraph('', style=self.configuration['style_mappings']['number'])
        self.handle_title(song)
        songparts = song.content
        for songpart in songparts[1:]:
            self.add_songpart_to_document(songpart)


    def handle_title(self, song):
        pattern = song.title
        # pattern = pattern.replace('.', '\.')
        # pattern = '^' + pattern.replace(' ', '*.')
        first_part = song.content[0]
        embedded_title = self.find_title_in_songpart(song.title, first_part)
        if embedded_title is None:
            self.add_title(song.title)
            self.add_songpart_to_document(first_part)
        else:
            self.add_songpart_including_title(first_part, embedded_title)

    def find_title_in_songpart(self, title, songpart):
        pattern = title
        match = re.search(pattern, songpart.lyrics())
        if match is None:
            return None
        return match.group()

    def add_title(self, title):
        style_name = self.configuration['style_mappings']['title']
        self.document.add_paragraph(title, style=style_name)

    def add_songpart_including_title(self, songpart, title):
        self.add_songpart_to_document(songpart)

    def add_songpart_to_document(self, songpart):
        self.previous_part_type = songpart.type
        style_name = self.configuration['style_mappings'][songpart.type]
        if not self.previous_indented:
            style_name += '-indent'
        self.previous_indented = not self.previous_indented
        self.document.add_paragraph(songpart.lyrics(), style=style_name)